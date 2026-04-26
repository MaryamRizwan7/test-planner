# venue_plan.py
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import _Cell
from datetime import datetime
import pandas as pd
from io import BytesIO

def get_day_with_suffix(day):
    """ Returns a day number with the correct suffix (e.g., 1st, 2nd, 24th). """
    if 4 <= day <= 20 or 24 <= day <= 30:
        return f"{day}th"
    else:
        # This will correctly handle 1, 2, 3, 21, 22, 23, 31
        # Example: day % 10 - 1 -> 1%10-1=0 ('st'), 2%10-1=1 ('nd'), 3%10-1=2 ('rd')
        # If result is not 0,1,2 (e.g., 4%10-1=3), it defaults to 'th'
        return f"{day}{['st', 'nd', 'rd', 'th'][min(day % 10 - 1, 3)]}"

def format_header_dates(date_objects):
    """
    Formats a list of dates into a single, compact header string.
    Example: [date(4), date(5), date(6)] -> "4th, 5th & 6th September 2025"
    """
    if not date_objects:
        return ""
    
    # Ensure date_objects are datetime.date or datetime.datetime objects
    sorted_dates = sorted(date_objects)
    
    # Get the month and year from the first date. Assumes all dates are in the same month/year.
    first_date = sorted_dates[0]
    month_year = first_date.strftime("%B %Y")
    
    # Get day numbers with suffixes
    day_numbers = [get_day_with_suffix(d.day) for d in sorted_dates]
    
    if len(day_numbers) > 2:
        # For more than 2 dates, join first n-1 with commas and last with " & "
        day_string = ", ".join(day_numbers[:-1]) + " & " + day_numbers[-1]
    elif len(day_numbers) == 2:
        # For exactly 2 dates, join with " & "
        day_string = " & ".join(day_numbers)
    else:
        # Single day
        day_string = day_numbers[0]
    
    return f"{day_string} {month_year}"

def format_table_date(dt):
    """ Formats a date for the table row. Example: "4th September 2025" """
    day_with_suffix = get_day_with_suffix(dt.day)
    return dt.strftime(f"{day_with_suffix} %B %Y")

def set_table_cell_font(cell: _Cell):
    """ Applies the specific font styling for all table cells. """
    cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.name = 'Calibri (Body)'
            run.font.size = Pt(14)
            run.bold = False

def convert_time_to_minutes(time_str):
    """ Convert time string (e.g., "09:00 AM") to minutes since midnight for sorting """
    if not time_str:
        return 0
    
    try:
        # Parse the time string
        time_obj = datetime.strptime(time_str, "%I:%M %p")
        # Convert to minutes since midnight
        return time_obj.hour * 60 + time_obj.minute
    except ValueError: 
        return 0 

def format_department_name(department_name):
    """ Format the department name for display """
    if "Civil" in department_name:
        return "Civil Engineering"
    elif "Computer" in department_name:
        return "Computer Science & Information Technology"
    elif "Urban" in department_name:
        return "Urban Engineering"
    else:
        return department_name  # Fallback for any other unexpected department names

def parse_date_string(date_str):
    """
    Parses a date string from MM-DD-YYYY format.
    This is designed to correctly interpret dates like "09-05-2025" as September 5th, 2025.
    """
    if not date_str:
        return None # Handle empty strings
    
    # First, try to parse with pandas using monthfirst=True
    try:
        return pd.to_datetime(date_str, monthfirst=True).to_pydatetime()
    except:
        pass
    
    # If that fails, try explicit formats, prioritizing MM-DD-YYYY
    formats_to_try = [
        "%m-%d-%Y",  # MM-DD-YYYY (e.g., 09-05-2025 as September 5th)
        "%m/%d/%Y",  # MM/DD/YYYY (e.g., 09/05/2025 as September 5th)
        "%d-%m-%Y",  # DD-MM-YYYY (e.g., 05-09-2025 as May 9th)
        "%d/%m/%Y",  # DD/MM/YYYY (e.g., 05/09/2025 as May 9th)
        "%Y-%m-%d",  # YYYY-MM-DD (e.g., 2025-09-05 as September 5th)
    ]
    
    for fmt in formats_to_try:
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue
    
    # As a last resort, try without specifying format
    try:
        return pd.to_datetime(date_str).to_pydatetime()
    except Exception:
        raise ValueError(f"Could not parse date string: '{date_str}'. Please ensure dates are in a recognizable format (e.g., MM-DD-YYYY, MM/DD/YYYY).")

def generate_venue_plan(schedule_df, program="master", semester="Fall", year=2025):
    """
    Generates a professionally formatted Word document with a 3-column table
    and dynamic multi-date headers.
    """
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    BLUE_COLOR = RGBColor(0, 0, 255)
    RED_COLOR = RGBColor(255, 0, 0)
    DARK_BLUE_COLOR = RGBColor(0, 0, 139)
    
    # Build column map to handle different case variations
    col_map = {c.lower(): c for c in schedule_df.columns}
    
    # Ensure 'day_object' is present and is datetime objects
    day_col = col_map.get('day')
    if not day_col:
        raise KeyError("Schedule data is missing 'day' column")
    
    # Parse dates using the consolidated function, now prioritizing MM-DD-YYYY
    schedule_df['day_object'] = schedule_df[day_col].apply(parse_date_string)
    
    unique_dates = sorted(schedule_df['day_object'].dt.date.unique()) 
    
    # Determine the header date string based on unique dates
    header_date_string = format_header_dates(unique_dates)
    
    # Get the day of the week for the *first* date if only one unique date
    day_of_week_prefix = ""
    if len(unique_dates) == 1:
        day_of_week_prefix = unique_dates[0].strftime('%A ')
    
    venue_col = col_map.get('venue')
    if not venue_col:
        raise KeyError("Schedule data is missing 'venue' column")
    
    # Group by venue column
    grouped = schedule_df.groupby(venue_col)
    
    docx_buffer = BytesIO()
    
    for block_name, rows in grouped:
        # Get the raw department name from the first row of the current block's data
        dept_col = col_map.get('department')
        if not dept_col:
            raise KeyError("Schedule data is missing 'department' column")
        
        raw_department_name = rows.iloc[0].get(dept_col, "")
        
        # Format the department name for display
        formatted_department_name = format_department_name(raw_department_name)
        
        # Add initial spacing
        for _ in range(8): doc.add_paragraph()
        
        # University Header
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("NED UNIVERSITY OF ENGINEERING & TECHNOLOGY")
        run.font.name = 'Balthazar'; run.font.size = Pt(20); run.bold = True
        p.paragraph_format.space_after = Pt(0)
        
        # Department Header (Dynamic) - Always show Computer Science & Information Technology
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Computer Science & Information Technology") 
        run.font.name = 'Balthazar'; run.font.size = Pt(18); run.bold = True  
        p.paragraph_format.space_after = Pt(6)
        
        # Program Header
        program_header = (
            f"Admission Entry Test BS CSIT/IS/DS and BSDS (GDS) ({semester}-{year})"
            if program.lower() == "bachelor" else 
            f"Admission Entry Test MS CSIT/IS/DS and MSDS (GDS) ({semester}-{year})"
        )
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(program_header)
        run.font.name = 'Arial'; run.font.size = Pt(16); run.bold = True
        run.font.color.rgb = BLUE_COLOR
        p.paragraph_format.space_after = Pt(0)
        
        # Admission Test Date
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Admission Test Date: {day_of_week_prefix}{header_date_string}")
        run.font.name = 'Arial'; run.font.size = Pt(19); run.bold = True
        run.font.color.rgb = RED_COLOR
        
        doc.add_paragraph()
        
        # Venue Header - Use the actual department name from the schedule
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Venue: "); run.font.name = 'Calibri Light (Headings)'; run.font.size = Pt(26); run.bold = True; run.font.color.rgb = DARK_BLUE_COLOR
        run2 = p.add_run(formatted_department_name) 
        run2.font.name = 'Calibri Light (Headings)'; run2.font.size = Pt(28); run2.bold = True; run2.font.color.rgb = DARK_BLUE_COLOR  
        
        doc.add_paragraph()
        
        # Block Name
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Block: {block_name}")
        run.font.name = 'Calibri (Body)'; run.font.size = Pt(36); run.bold = True
        p.paragraph_format.space_after = Pt(18)
        
        # Table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Date"
        hdr_cells[1].text = "Time"
        hdr_cells[2].text = "Roll No"
        for cell in hdr_cells:
            set_table_cell_font(cell)
            for p_cell in cell.paragraphs:
                for run_cell in p_cell.runs:
                    run_cell.bold = True
        
        # Get time and roll number columns
        time_col = col_map.get('time')
        roll_col = col_map.get('roll no') or col_map.get('roll_no') or col_map.get('roll')
        
        if not time_col:
            raise KeyError("Schedule data is missing 'time' column")
        if not roll_col:
            raise KeyError("Schedule data is missing 'roll no' column")
        
        # Create a copy of the rows for sorting
        sorted_rows = rows.copy()
        
        # Convert time strings to minutes for proper sorting
        sorted_rows['time_minutes'] = sorted_rows[time_col].apply(convert_time_to_minutes)
        
        # Sort by date_object and then by time_minutes
        sorted_rows = sorted_rows.sort_values(by=['day_object', 'time_minutes'])
        
        # Add rows to the table in the correct order
        for _, row in sorted_rows.iterrows():
            row_cells = table.add_row().cells
            table_date_obj = row['day_object']
            row_cells[0].text = format_table_date(table_date_obj)
            row_cells[1].text = row.get(time_col, "")
            row_cells[2].text = row.get(roll_col, "")
            for cell in row_cells:
                set_table_cell_font(cell)
        
        table.columns[0].width = Inches(2.5)
        table.columns[1].width = Inches(2.5)
        table.columns[2].width = Inches(2.5)
        
        doc.add_page_break()
    
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

# -------- NEW: Preview function for frontend --------
def generate_venue_plan_preview(schedule_df, program="master"):
    """
    Generate a plain-text preview version of the venue plan (for React display).
    Returns a string that can be split and shown in blocks.
    """
    preview_blocks = []
    
    # Build column map to handle different case variations
    col_map = {c.lower(): c for c in schedule_df.columns}
    
    # Ensure 'day_object' is present and is datetime objects
    day_col = col_map.get('day')
    if not day_col:
        raise KeyError("Schedule data is missing 'day' column")
    
    # Parse dates using the consolidated function
    schedule_df['day_object'] = schedule_df[day_col].apply(parse_date_string)
    
    venue_col = col_map.get('venue')
    if not venue_col:
        raise KeyError("Schedule data is missing 'venue' column")
    
    grouped = schedule_df.groupby(venue_col) # Group by venue column
    
    # Get department, incharge, time, and roll number columns
    dept_col = col_map.get('department')
    incharge_col = col_map.get('department incharge') or col_map.get('incharge')
    time_col = col_map.get('time')
    roll_col = col_map.get('roll no') or col_map.get('roll_no') or col_map.get('roll')
    
    if not dept_col:
        raise KeyError("Schedule data is missing 'department' column")
    if not incharge_col:
        raise KeyError("Schedule data is missing 'incharge' column")
    if not time_col:
        raise KeyError("Schedule data is missing 'time' column")
    if not roll_col:
        raise KeyError("Schedule data is missing 'roll no' column")
    
    for block_name, rows in grouped:
        # Get the raw department name from the first row of the current block's data
        raw_department_name = rows.iloc[0].get(dept_col, "")
        
        # Format the department name for display
        formatted_department_name = format_department_name(raw_department_name)
        
        # Create a copy of the rows for sorting
        sorted_rows = rows.copy()
        
        # Convert time strings to minutes for proper sorting
        sorted_rows['time_minutes'] = sorted_rows[time_col].apply(convert_time_to_minutes)
        
        # Sort by date_object and then by time_minutes
        sorted_rows = sorted_rows.sort_values(by=['day_object', 'time_minutes'])
        
        incharge = rows.iloc[0].get(incharge_col, "N/A")
        
        block_lines = [
            f"--- VENUE: Block {block_name} ({formatted_department_name}) ---",  # Use the formatted department name
            f"Incharge: {incharge}",
            "-" * 50,
            f"{'Date':<20} | {'Time':<15} | {'Roll No':<10}",
            "-" * 50
        ]
        
        # Add rows to the block lines in the correct order
        for _, row in sorted_rows.iterrows():
            date_str = format_table_date(row['day_object'])
            time_str = row.get(time_col, "")
            roll_str = row.get(roll_col, "")
            block_lines.append(f"{date_str:<20} | {time_str:<15} | {roll_str:<10}")
        
        preview_blocks.append("\n".join(block_lines))
    
    return "\n\n".join(preview_blocks)